import requests #type:ignore
from bs4 import BeautifulSoup #type:ignore
from urllib.parse import urljoin, urlparse
from langchain_core.documents import Document #type:ignore
from typing import List, Union
import io
import PyPDF2  # For PDF processing #type:ignore
import docx #type:ignore
import docx2txt #type:ignore
import pypandoc #type:ignore
from typing import List

def parse_docx(uploaded_file) -> List[Document]:
    """Extract text from DOCX using python-docx with page breaks"""
    documents = []
    try:
        # Option 1: Using python-docx (better formatting)
        doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
        full_text = []
        
        for para in doc.paragraphs:
            # Detect page breaks (WARNING: imperfect detection)
            if para.runs and para.runs[0].text == '\x0c':
                if full_text:  # Save current page
                    documents.append(create_docx_document(uploaded_file, len(documents)+1, '\n'.join(full_text)))
                    full_text = []
            else:
                full_text.append(para.text)
        
        if full_text:  # Add remaining content
            documents.append(create_docx_document(uploaded_file, len(documents)+1, '\n'.join(full_text)))
            
        # Fallback if no content extracted (shouldn't happen with docx)
        if not documents:
            text = docx2txt.process(io.BytesIO(uploaded_file.getvalue()))
            if text.strip():
                documents.append(create_docx_document(uploaded_file, 1, text))
                
    except Exception as e:
        # Fallback to pypandoc if other methods fail
        try:
            text = pypandoc.convert_text(uploaded_file.getvalue(), 'plain', format='docx')
            if text.strip():
                documents.append(create_docx_document(uploaded_file, 1, text))
        except:
            raise Exception(f"DOCX processing error: {str(e)}")
    
    return documents

def create_docx_document(uploaded_file, page_num: int, text: str) -> Document:
    """Helper to create Document with metadata"""
    return Document(
        page_content=text,
        metadata={
            "source": uploaded_file.name,
            "page": page_num,
            "type": "docx"
        }
    )

# For older .doc files (requires pypandoc)
def parse_doc(uploaded_file) -> List[Document]:
    """Process older .doc format files"""
    try:
        text = pypandoc.convert_text(uploaded_file.getvalue(), 'plain', format='doc')
        if text.strip():
            return [
                Document(
                    page_content=text,
                    metadata={
                        "source": uploaded_file.name,
                        "page": 1,
                        "type": "doc"
                    }
                )
            ]
        return []
    except Exception as e:
        raise Exception(f"DOC processing error: {str(e)}")

def scrape_website(url: str, max_pages: int = 3) -> List[Document]:
    """Scrape content from a website."""
    if not is_valid_url(url):
        raise ValueError(f"Invalid URL: {url}")
    
    base_url = get_base_url(url)
    visited_urls = set()
    documents = []
    
    try:
        queue = [url]
        
        while queue and len(visited_urls) < max_pages:
            current_url = queue.pop(0)
            
            if current_url in visited_urls:
                continue
                
            try:
                headers = {'User-Agent': 'Mozilla/5.0'}
                response = requests.get(current_url, headers=headers, timeout=10)
                response.raise_for_status()
                
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Remove unwanted elements
                for element in soup(['script', 'style', 'nav', 'footer', 'iframe']):
                    element.decompose()
                
                title = soup.title.string if soup.title else "No Title"
                main_content = soup.find('main') or soup.find('article') or soup.body
                text = main_content.get_text(separator='\n', strip=True) if main_content else ""
                
                if text:
                    documents.append(
                        Document(
                            page_content=f"URL: {current_url}\nTitle: {title}\nContent: {text}",
                            metadata={"source": current_url, "type": "webpage"}
                        )
                    )
                
                visited_urls.add(current_url)
                
                # Find and add new links to queue
                if len(visited_urls) < max_pages:
                    for link in soup.find_all('a', href=True):
                        href = link['href']
                        absolute_url = urljoin(base_url, href)
                        if (absolute_url.startswith(base_url) and 
                            absolute_url not in visited_urls and 
                            absolute_url not in queue):
                            queue.append(absolute_url)
                
            except Exception as e:
                print(f"Error scraping {current_url}: {str(e)}")
                continue
                
    except Exception as e:
        raise Exception(f"Scraping failed: {str(e)}")
    
    return documents

def parse_pdf(uploaded_file) -> List[Document]:
    """Extract text from PDF with page-wise chunking."""
    documents = []
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.getvalue()))
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            
            if text.strip():
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": uploaded_file.name,
                            "page": page_num + 1,
                            "type": "pdf"
                        }
                    )
                )
    except Exception as e:
        raise Exception(f"PDF processing error: {str(e)}")
    
    return documents

def is_valid_url(url: str) -> bool:
    """Check if URL is valid."""
    parsed = urlparse(url)
    return bool(parsed.netloc) and bool(parsed.scheme)

def get_base_url(url: str) -> str:
    """Extract base URL."""
    parsed = urlparse(url)
    return f"{parsed.scheme}://{parsed.netloc}"